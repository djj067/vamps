"""
Meeting recorder demo backend.

Flow:
  browser records mic in short segments
    -> POST /transcribe  (each segment -> Deepgram transcription -> text)
    -> browser appends text to a live transcript
    -> POST /spec        (full transcript -> LLM -> functional spec)

Everything is kept deliberately thin so the pipeline is easy to see and debug.
"""
import asyncio
import io
import json
import os
import re
import shutil
import subprocess
import threading
import time
import uuid
from pathlib import Path

import httpx
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, UploadFile
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openai import OpenAI

load_dotenv()

# --- Transcription: Deepgram -------------------------------------------------
DEEPGRAM_API_KEY = os.getenv("DEEPGRAM_API_KEY", "")
DEEPGRAM_URL = "https://api.deepgram.com/v1/listen"
# nova-3 is Deepgram's latest/most accurate general model.
TRANSCRIBE_MODEL = os.getenv("TRANSCRIBE_MODEL", "nova-3")
# Force English so short segments aren't misdetected as another language.
TRANSCRIBE_LANGUAGE = os.getenv("TRANSCRIBE_LANGUAGE", "en")
# Comma-separated domain terms (names, product words, Singlish) to bias toward.
# Deepgram calls this "keyterm prompting" on nova-3.
TRANSCRIBE_KEYTERMS = [
    t.strip() for t in os.getenv("TRANSCRIBE_KEYTERMS", "").split(",") if t.strip()
]

# --- Spec generation: Claude (via the CLI for testing, OpenRouter for demo) --
# Provider is runtime-switchable from the UI (POST /provider) — no env editing.
# SPEC_PROVIDER is just the startup default.
SPEC_PROVIDER = os.getenv("SPEC_PROVIDER", "claude_cli").lower()
RUNTIME = {"provider": SPEC_PROVIDER}


def current_provider() -> str:
    return RUNTIME["provider"]
SPEC_MODEL = os.getenv("SPEC_MODEL", "anthropic/claude-3.5-sonnet")   # openrouter slug
SPEC_MODEL_CLI = os.getenv("SPEC_MODEL_CLI", "")                       # optional --model for the CLI
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_BASE_URL = os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")

# --- Langflow bridge ---------------------------------------------------------
# When LANGFLOW_API_KEY is set, /spec and /build run through Langflow flows
# instead of the inline logic below. /transcribe always stays direct (low latency
# for live mic segments). Flows are resolved by display name via the Langflow API.
LANGFLOW_URL = os.getenv("LANGFLOW_URL", "http://127.0.0.1:7860").rstrip("/")
# Optional — only needed if the Langflow server enforces auth. For a local server
# with auth disabled, leave it blank.
LANGFLOW_API_KEY = os.getenv("LANGFLOW_API_KEY", "")
LANGFLOW_SPEC_FLOW = os.getenv("LANGFLOW_SPEC_FLOW", "VAMPS Spec API")
LANGFLOW_BUILD_FLOW = os.getenv("LANGFLOW_BUILD_FLOW", "VAMPS Build API")
# The node ids inside the API flows (generate_flows.py) that tweaks target.
LANGFLOW_SPEC_NODE = os.getenv("LANGFLOW_SPEC_NODE", "Spec-vsa")
LANGFLOW_BUILD_NODE = os.getenv("LANGFLOW_BUILD_NODE", "Build-vba")
# Claude Code can take several minutes on a full prototype; give it room.
LANGFLOW_BUILD_TIMEOUT = int(os.getenv("LANGFLOW_BUILD_TIMEOUT", "900"))  # seconds
# Flip the bridge on with LANGFLOW_ENABLED=true (no API key required locally).
USE_LANGFLOW = os.getenv("LANGFLOW_ENABLED", "false").strip().lower() in ("1", "true", "yes", "on")

BASE_DIR = Path(__file__).parent
REC_DIR = BASE_DIR / "recordings"
REC_DIR.mkdir(exist_ok=True)
GEN_DIR = BASE_DIR / "generated"          # Claude Code writes prototypes here
GEN_DIR.mkdir(exist_ok=True)

# Claude Code CLI — runs headless to generate the prototype. Authenticates via
# CLAUDE_CODE_OAUTH_TOKEN (inherited from the environment) against the CC plan.
CLAUDE_BIN = shutil.which("claude") or "claude"

app = FastAPI(title="VAMPS")

# In-memory registry of prototype build jobs: id -> {proc, dir, started}
BUILD_JOBS: dict = {}
# Prototype "projects": project_id -> {"dir": Path, "session_id": str|None}. Lets a
# rebuild resume the same Claude Code session and edit the existing index.html.
PROJECTS: dict = {}


# The model returns a JSON object: {"spec": "<markdown>", "questions": [ ... ]}.
# `spec` follows this structure; `questions` are the analyst's live clarifications.
SPEC_STRUCTURE = """# Functional Specification

## 1. Meeting Summary
2-4 sentences: who/what/why.

## 2. Client Goals & Pain Points
Bulleted list.

## 3. Proposed Solution / Scope
What we will build or deliver.

## 4. Functional Requirements
Numbered list. Each item: a concrete, testable requirement.

## 5. Non-Functional Requirements
Performance, security, compliance, etc. (only if mentioned or clearly implied.)

## 6. Action Items
Who does what next (use "Unassigned" if the transcript doesn't say)."""

SPEC_SYSTEM = f"""You are a business analyst supporting a sales team. You turn
the raw material of a client meeting (a transcript that may contain
transcription errors, filler words and overlapping speakers, plus any
pre-existing notes) into a clear FUNCTIONAL SPECIFICATION.

FIRST, filter the material. A real meeting transcript is full of content that is
NOT about what to build: greetings and small talk, scheduling and logistics,
side conversations, jokes, weather, off-topic tangents, people testing the mic,
and unrelated chit-chat. IGNORE all of that. Base the specification ONLY on
content relevant to the client's product, goals, needs, scope, and requirements.
Never turn an off-topic remark into a requirement, an action item, or a
clarifying question. If a whole segment is irrelevant, leave the spec unchanged.

You ALWAYS respond with a single JSON object with exactly these keys:
  "spec":      a Markdown string using EXACTLY this structure and headings:
{SPEC_STRUCTURE}
  "questions": an array of objects {{"question": <string>, "anchor": <string>}}
               — the clarifying questions you, as the analyst, STILL need
               answered (things the material leaves ambiguous). For each:
                 - "question": one specific question, one sentence.
                 - "anchor": a SHORT verbatim quote (a few words) copied EXACTLY
                   from your "spec" markdown identifying the sentence or
                   requirement the question is about, so it can be shown next to
                   that text. Copy it word-for-word from the spec; if the
                   question is general, use "".
               Do NOT include a question already answered by the material.
               Return [] if nothing needs clarifying.
  "answered":  an array of objects {{"question": <string>, "answer": <string>,
               "evidence": <string>}}. Include a CURRENTLY OPEN QUESTION here
               ONLY if the material EXPLICITLY answers it:
                 - "question": the open question's exact text.
                 - "answer": one sentence stating the answer.
                 - "evidence": a SHORT quote copied WORD-FOR-WORD from the
                   material that states the answer. Do NOT paraphrase, translate,
                   summarise, or invent this quote — copy it verbatim.
               A topic merely being MENTIONED is NOT the same as being answered.
               If you would have to guess, infer, assume, or read between the
               lines, DO NOT put it here — leave it in "questions". When in
               doubt, leave the question open. Return [] if nothing is
               explicitly resolved this turn.

Be concrete. If something was not discussed, write "Not discussed" rather than
inventing details. Do NOT put a questions section inside the spec markdown;
clarifications belong only in "questions" / "answered"."""

# First pass: no prior spec exists yet.
SPEC_USER_INITIAL = """Produce the functional specification from the material below.

MATERIAL:
---
{transcript}
---"""

# Later passes: revise the existing spec in light of newly transcribed material.
SPEC_USER_REVISE = """A functional specification already exists (below). NEW
material has since been transcribed. REVISE and extend the existing spec so it
reflects everything now known — keep what still holds, correct what changed, and
fold in the new details. Do not discard prior content just because it is not
repeated in the new material.

{resolved_block}{open_block}EXISTING SPECIFICATION:
---
{previous_spec}
---

FULL MATERIAL SO FAR (includes the new transcript):
---
{transcript}
---"""


VERIFY_SYSTEM = """You are a strict fact-checker for a business analyst. You are
given a meeting transcript and candidate items, each with a QUESTION, a proposed
ANSWER, and an EVIDENCE quote. For each, decide whether the transcript
EXPLICITLY and UNAMBIGUOUSLY answers the question.

Mark answered = false when ANY of these hold:
  - the topic is only mentioned in passing;
  - the speaker defers or hedges it ("decide later", "not sure", "maybe", "we'll see");
  - answering would require guessing, inference, or assumption;
  - the evidence quote does not, on its own, directly state the answer.

Be conservative: if there is any doubt, answered = false.

Respond with a single JSON object:
{"results": [{"question": "<the exact question text>", "answered": true or false}]}"""


def _usage_of(resp) -> dict:
    """Extract OpenAI token usage from a chat-completions response, defensively."""
    u = getattr(resp, "usage", None)
    return {
        "prompt_tokens": getattr(u, "prompt_tokens", 0) or 0,
        "completion_tokens": getattr(u, "completion_tokens", 0) or 0,
        "total_tokens": getattr(u, "total_tokens", 0) or 0,
    }


def _parse_json(text: str) -> dict:
    """Parse a JSON object from an LLM reply, tolerating code fences / preamble."""
    t = (text or "").strip()
    if t.startswith("```"):
        t = re.sub(r"^```[a-zA-Z]*\s*", "", t)
        t = re.sub(r"\s*```$", "", t)
    try:
        return json.loads(t)
    except Exception:
        i, j = t.find("{"), t.rfind("}")
        if i >= 0 and j > i:
            return json.loads(t[i:j + 1])
        raise


def _spec_complete(system: str, user: str, temperature: float = 0.3):
    """One Claude completion, returning (content_text, usage_dict).

    claude_cli: local Claude auth via the CLI (no API cost) — for testing.
    openrouter: OpenAI-compatible call to OpenRouter — for the demo.
    """
    if current_provider() == "openrouter":
        resp = OpenAI(base_url=OPENROUTER_BASE_URL, api_key=OPENROUTER_API_KEY).chat.completions.create(
            model=SPEC_MODEL, temperature=temperature,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        )
        return (resp.choices[0].message.content or ""), _usage_of(resp)

    cmd = [CLAUDE_BIN, "-p", "--output-format", "json", "--allowed-tools", ""]
    if SPEC_MODEL_CLI:
        cmd += ["--model", SPEC_MODEL_CLI]
    prompt = f"{system}\n\n{user}"
    wrapper, last = {}, ""
    for _ in range(2):   # CLI occasionally returns empty stdout; retry once
        proc = subprocess.run(cmd, input=prompt, capture_output=True, text=True, timeout=240)
        last = proc.stderr or ""
        if (proc.stdout or "").strip():
            wrapper = json.loads(proc.stdout)
            if wrapper.get("result"):
                break
    result = wrapper.get("result") or ""
    if not result:
        raise RuntimeError(f"claude CLI returned no output. stderr: {last[:200]}")
    u = wrapper.get("usage") or {}
    inp = (u.get("input_tokens", 0) or 0) + (u.get("cache_read_input_tokens", 0) or 0) \
        + (u.get("cache_creation_input_tokens", 0) or 0)
    out = u.get("output_tokens", 0) or 0
    return result, {"prompt_tokens": inp, "completion_tokens": out, "total_tokens": inp + out}


def _verify_answers(candidates: list, transcript: str):
    """Second-pass skeptic: keep only candidates the model confirms are truly answered.

    Returns (kept_list, usage_dict) so the caller can tally verification tokens.
    """
    zero = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
    if not candidates:
        return [], zero
    listing = "\n\n".join(
        f'{i+1}. QUESTION: {c["question"]}\n   PROPOSED ANSWER: {c["answer"]}\n   EVIDENCE: "{c["evidence"]}"'
        for i, c in enumerate(candidates)
    )
    user = f"TRANSCRIPT:\n---\n{transcript}\n---\n\nCANDIDATES:\n{listing}"
    try:
        content, usage = _spec_complete(VERIFY_SYSTEM, user, temperature=0)
        data = _parse_json(content)
        confirmed = {
            _normalize(str(r.get("question", "")))
            for r in data.get("results", [])
            if r.get("answered") is True
        }
        return [c for c in candidates if _normalize(c["question"]) in confirmed], usage
    except Exception:
        # On verifier failure, fall back to the evidence-checked candidates rather than crash.
        return candidates, zero


def _normalize(s: str) -> str:
    """Lowercase, strip punctuation, collapse whitespace — for grounding checks."""
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9 ]+", " ", s.lower())).strip()


def _evidence_supported(evidence: str, transcript: str) -> bool:
    """True only if the model's quote is actually present in the material.

    Guards against the model marking a question "answered" when the conversation
    merely touched the topic. Requires either a verbatim substring match, or a
    high overlap of the quote's significant words (tolerates minor paraphrase).
    """
    ev = _normalize(evidence)
    if len(ev) < 4:
        return False
    tx = _normalize(transcript)
    if ev in tx:
        return True
    words = {w for w in ev.split() if len(w) > 3}
    if not words:
        return False
    hits = sum(1 for w in words if w in tx)
    return hits / len(words) >= 0.8


def _build_spec_messages(
    transcript: str, previous_spec: str, resolved: list, open_questions: list
) -> list:
    """Assemble the chat messages for either a first-pass or a revision call."""
    if previous_spec.strip():
        resolved_block = ""
        if resolved:
            joined = "\n".join(f"- {q}" for q in resolved)
            resolved_block = (
                "The following questions were already resolved by the client — "
                "treat them as answered and do NOT raise them again:\n"
                f"{joined}\n\n"
            )
        open_block = ""
        if open_questions:
            joined = "\n".join(f"- {q}" for q in open_questions)
            open_block = (
                "CURRENTLY OPEN QUESTIONS (still awaiting an answer). If the "
                "latest material now answers any of these, move it into "
                '"answered" with a one-sentence answer instead of repeating it:\n'
                f"{joined}\n\n"
            )
        user = SPEC_USER_REVISE.format(
            resolved_block=resolved_block,
            open_block=open_block,
            previous_spec=previous_spec,
            transcript=transcript,
        )
    else:
        user = SPEC_USER_INITIAL.format(transcript=transcript)
    return [
        {"role": "system", "content": SPEC_SYSTEM},
        {"role": "user", "content": user},
    ]


# --- Langflow client ---------------------------------------------------------
_FLOW_ID_CACHE: dict = {}


def _lf_headers() -> dict:
    h = {"Content-Type": "application/json"}
    if LANGFLOW_API_KEY:            # only sent when a key is configured
        h["x-api-key"] = LANGFLOW_API_KEY
    return h


# Flow name -> JSON file main.py (re)imports if the flow is missing or its edges
# were pruned. Langflow drops non-native edges on restart, so we self-heal.
_FLOW_FILES = {
    LANGFLOW_SPEC_FLOW: BASE_DIR / "langflow_components" / "vamps_spec_api_flow.json",
    LANGFLOW_BUILD_FLOW: BASE_DIR / "langflow_components" / "vamps_build_api_flow.json",
}


def _lf_find_id(name: str):
    with httpx.Client(timeout=30) as http:
        resp = http.get(f"{LANGFLOW_URL}/api/v1/flows/",
                        params={"get_all": "true", "header_flows": "true"},
                        headers=_lf_headers())
    resp.raise_for_status()
    body = resp.json()
    flows = body.get("items", body) if isinstance(body, dict) else body
    for f in flows:
        if f.get("name") == name:
            return f["id"]
    return None


def _lf_edge_count(flow_id: str) -> int:
    with httpx.Client(timeout=30) as http:
        resp = http.get(f"{LANGFLOW_URL}/api/v1/flows/{flow_id}", headers=_lf_headers())
    if resp.status_code != 200:
        return -1
    return len(((resp.json() or {}).get("data") or {}).get("edges", []))


def _lf_import(flow_file: Path) -> str:
    with httpx.Client(timeout=30) as http:
        resp = http.post(f"{LANGFLOW_URL}/api/v1/flows/",
                         headers=_lf_headers(), content=flow_file.read_bytes())
    resp.raise_for_status()
    return resp.json()["id"]


def _lf_delete(flow_id: str) -> None:
    with httpx.Client(timeout=30) as http:
        http.delete(f"{LANGFLOW_URL}/api/v1/flows/{flow_id}", headers=_lf_headers())


def _lf_flow_id(name: str) -> str:
    """Resolve the flow id, self-healing if the flow is missing or edge-pruned.

    Langflow drops non-native edges on restart ("graph has vertices but no
    edges") and the id can go stale if a flow is deleted. `_lf_edge_count`
    returns >0 (healthy), 0 (edges pruned), or -1 (flow gone). Anything not >0
    means rebuild from the JSON, which restores a healthy flow.
    """
    flow_file = _FLOW_FILES.get(name)

    cached = _FLOW_ID_CACHE.get(name)
    if cached and _lf_edge_count(cached) > 0:
        return cached                       # cached and healthy

    found = _lf_find_id(name)               # re-resolve by name
    if found and _lf_edge_count(found) > 0:
        _FLOW_ID_CACHE[name] = found
        return found

    if found:                               # exists but edgeless -> drop it
        _lf_delete(found)
    if not flow_file or not flow_file.exists():
        raise RuntimeError(f"Langflow flow '{name}' not found and no JSON to import.")
    fid = _lf_import(flow_file)             # rebuild from JSON (restores edges)
    _FLOW_ID_CACHE[name] = fid
    return fid


def _extract_run_json(resp_json: dict) -> dict:
    """Pull the JSON string our Text Output emitted out of the /run envelope.

    The envelope nests component results in different shapes across versions, so
    we collect every candidate string and return the first that parses to a dict.
    """
    candidates = []

    def walk(o):
        if isinstance(o, str):
            candidates.append(o)
        elif isinstance(o, dict):
            for v in o.values():
                walk(v)
        elif isinstance(o, list):
            for v in o:
                walk(v)

    for out in resp_json.get("outputs", []):
        walk(out.get("outputs", out))
    for s in candidates:
        s = s.strip()
        if s.startswith("{"):
            try:
                d = json.loads(s)
                if isinstance(d, dict) and ("spec" in d or "html" in d or "status" in d):
                    return d
            except Exception:
                continue
    raise RuntimeError("Could not find a JSON result in the Langflow response.")


def _lf_run(name: str, input_value: str, tweaks: dict, timeout: float = 300) -> dict:
    """POST to /api/v1/run/{flow_id} and return the parsed JSON result. Blocking."""
    flow_id = _lf_flow_id(name)
    payload = {"input_value": input_value, "input_type": "chat",
               "output_type": "text", "tweaks": tweaks or {}}
    with httpx.Client(timeout=timeout) as http:
        resp = http.post(f"{LANGFLOW_URL}/api/v1/run/{flow_id}",
                         headers=_lf_headers(), json=payload)
    resp.raise_for_status()
    return _extract_run_json(resp.json())


@app.get("/")
def index():
    # no-store so an edited UI always loads fresh (this app is under active iteration)
    return FileResponse(
        BASE_DIR / "static" / "index.html",
        headers={"Cache-Control": "no-store, must-revalidate"},
    )


@app.post("/provider")
async def set_provider(payload: dict):
    """Switch spec-gen provider at runtime (claude_cli | openrouter) — no env edit."""
    p = str(payload.get("provider", "")).strip().lower()
    if p not in ("claude_cli", "openrouter"):
        return JSONResponse(status_code=400, content={"ok": False, "error": "provider must be claude_cli or openrouter"})
    RUNTIME["provider"] = p
    return {"ok": True, "provider": p}


@app.get("/usage")
async def usage_info():
    """Which model spec-gen uses, plus live OpenRouter credit when in demo mode.

    OpenRouter's /api/v1/key returns cumulative `usage` (credits spent) and
    `limit_remaining`; the UI diffs `usage` across calls to show per-run spend.
    """
    openrouter = current_provider() == "openrouter"
    info = {
        "provider": current_provider(),
        "spec_label": "Claude (CLI)" if not openrouter else "OpenRouter",
        "spec_model": ("(CLI default)" if not (openrouter or SPEC_MODEL_CLI) else
                       (SPEC_MODEL if openrouter else SPEC_MODEL_CLI)),
        "build_label": "Claude Code",
    }
    if openrouter and OPENROUTER_API_KEY:
        try:
            async with httpx.AsyncClient(timeout=15) as http:
                r = await http.get("https://openrouter.ai/api/v1/key",
                                   headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}"})
            r.raise_for_status()
            d = r.json().get("data", {}) or {}
            info["openrouter"] = {
                "usage": d.get("usage"),
                "limit": d.get("limit"),
                "limit_remaining": d.get("limit_remaining"),
                "is_free_tier": d.get("is_free_tier"),
            }
        except Exception as e:
            info["openrouter_error"] = str(e)
    return info


@app.post("/transcribe")
async def transcribe(
    audio: UploadFile = File(...),
    seq: int = Form(0),
    session: str = Form("default"),
):
    """Transcribe one audio segment via Deepgram. Returns text + timing for the debug panel."""
    t0 = time.time()
    data = await audio.read()
    size_kb = round(len(data) / 1024, 1)

    # Keep a local backup of every segment (client meetings are high-stakes).
    fname = REC_DIR / f"{session}_{seq:04d}.webm"
    fname.write_bytes(data)

    if not DEEPGRAM_API_KEY:
        return JSONResponse(
            status_code=500,
            content={"ok": False, "seq": seq, "size_kb": size_kb,
                     "error": "DEEPGRAM_API_KEY is not set in .env"},
        )

    # Deepgram pre-recorded API: POST the raw audio bytes, options go in the query string.
    # smart_format adds punctuation/capitalization; language pins it to English.
    params = {
        "model": TRANSCRIBE_MODEL,
        "language": TRANSCRIBE_LANGUAGE,
        "smart_format": "true",
    }
    if TRANSCRIBE_KEYTERMS:
        params["keyterm"] = TRANSCRIBE_KEYTERMS  # repeated query param, one per term
    content_type = audio.content_type or "audio/webm"

    try:
        async with httpx.AsyncClient(timeout=60) as http:
            resp = await http.post(
                DEEPGRAM_URL,
                params=params,
                content=data,
                headers={
                    "Authorization": f"Token {DEEPGRAM_API_KEY}",
                    "Content-Type": content_type,
                },
            )
        resp.raise_for_status()
        body = resp.json()
        # transcript lives at results.channels[0].alternatives[0].transcript
        alt = body["results"]["channels"][0]["alternatives"][0]
        text = (alt.get("transcript") or "").strip()
        return {
            "ok": True,
            "seq": seq,
            "text": text,
            "size_kb": size_kb,
            "latency_s": round(time.time() - t0, 2),
            "model": TRANSCRIBE_MODEL,
            "confidence": alt.get("confidence"),
        }
    except httpx.HTTPStatusError as e:
        return JSONResponse(
            status_code=500,
            content={"ok": False, "seq": seq, "size_kb": size_kb,
                     "error": f"Deepgram {e.response.status_code}: {e.response.text[:300]}"},
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"ok": False, "seq": seq, "size_kb": size_kb, "error": str(e)},
        )


@app.post("/spec")
async def spec(payload: dict):
    """Generate or revise the functional spec, plus clarifying questions.

    First call (no previous_spec) writes the spec from scratch; every later call
    revises the existing spec in light of the latest transcript, so the model
    always builds on what it produced before.
    """
    transcript = (payload.get("transcript") or "").strip()
    previous_spec = (payload.get("previous_spec") or "").strip()
    resolved = [str(q).strip() for q in (payload.get("resolved_questions") or []) if str(q).strip()]
    open_qs = [str(q).strip() for q in (payload.get("open_questions") or []) if str(q).strip()]
    if not transcript:
        return JSONResponse(status_code=400, content={"ok": False, "error": "empty transcript"})

    # Langflow path: run the spec flow and pass it straight back to the browser.
    if USE_LANGFLOW:
        t0 = time.time()
        try:
            # Pass the runtime provider (+ model/key for openrouter) so the toggle
            # takes effect inside Langflow without touching its env.
            node_tweaks = {
                "previous_spec": previous_spec,
                "resolved_questions": "\n".join(resolved),
                "open_questions": "\n".join(open_qs),
                "provider": current_provider(),
            }
            if current_provider() == "openrouter":
                node_tweaks["model"] = SPEC_MODEL
                if OPENROUTER_API_KEY:
                    node_tweaks["api_key"] = OPENROUTER_API_KEY
            tweaks = {LANGFLOW_SPEC_NODE: node_tweaks}
            data = await asyncio.to_thread(_lf_run, LANGFLOW_SPEC_FLOW, transcript, tweaks)
            data["latency_s"] = round(time.time() - t0, 2)
            data["via"] = "langflow"
            return data
        except Exception as e:
            return JSONResponse(status_code=502, content={"ok": False, "error": f"Langflow spec failed: {e}"})

    t0 = time.time()
    try:
        msgs = _build_spec_messages(transcript, previous_spec, resolved, open_qs)
        content, usage = _spec_complete(msgs[0]["content"], msgs[1]["content"], temperature=0.3)
        data = _parse_json(content)   # tally spec-generation tokens (+ verify pass below)
        spec_md = (data.get("spec") or "").strip()
        # Questions are objects {question, anchor}; anchor is a quote used to place
        # the question next to the spec text it's about.
        questions = []
        for q in (data.get("questions") or []):
            if isinstance(q, dict):
                text, anchor = str(q.get("question", "")).strip(), str(q.get("anchor", "")).strip()
            else:
                text, anchor = str(q).strip(), ""
            if text:
                questions.append({"question": text, "anchor": anchor})
        # Auto-resolve a question only if (1) its verbatim evidence is really in the
        # material, and (2) a strict skeptic pass confirms it's genuinely answered.
        candidates = []
        for a in (data.get("answered") or []):
            if not isinstance(a, dict):
                continue
            q = str(a.get("question", "")).strip()
            if q and _evidence_supported(str(a.get("evidence", "")), transcript):
                candidates.append({
                    "question": q,
                    "answer": str(a.get("answer", "")).strip(),
                    "evidence": str(a.get("evidence", "")).strip(),
                })
        confirmed, verify_usage = _verify_answers(candidates, transcript)
        for k in usage:
            usage[k] += verify_usage.get(k, 0)
        answered = [{"question": c["question"], "answer": c["answer"]} for c in confirmed]
        # Any candidate that failed either gate stays visible as an open question.
        answered_texts = {a["question"] for a in answered}
        open_texts = {q["question"] for q in questions}
        for a in (data.get("answered") or []):
            if isinstance(a, dict):
                q = str(a.get("question", "")).strip()
                if q and q not in answered_texts and q not in open_texts:
                    questions.append({"question": q, "anchor": ""})
                    open_texts.add(q)
        questions = [q for q in questions if q["question"] not in answered_texts]
        return {
            "ok": True,
            "spec": spec_md,
            "questions": questions,
            "answered": answered,
            "revised": bool(previous_spec),
            "latency_s": round(time.time() - t0, 2),
            "model": SPEC_MODEL,
            "transcript_chars": len(transcript),
            "usage": usage,
        }
    except Exception as e:
        return JSONResponse(status_code=500, content={"ok": False, "error": str(e)})


@app.post("/upload_notes")
async def upload_notes(file: UploadFile = File(...)):
    """Ingest pre-existing meeting notes / a spec (.docx, .txt, .md) as seed material."""
    data = await file.read()
    name = file.filename or "notes"
    ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""
    try:
        if ext == "docx":
            doc = Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            # txt / md / anything text-like
            text = data.decode("utf-8", errors="replace")
    except Exception as e:
        return JSONResponse(status_code=400, content={"ok": False, "error": f"could not read {name}: {e}"})

    text = text.strip()
    if not text:
        return JSONResponse(status_code=400, content={"ok": False, "error": "file had no readable text"})
    return {"ok": True, "name": name, "text": text, "chars": len(text)}


def _add_runs(paragraph, text):
    """Render inline **bold** markdown into docx runs."""
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1  # odd segments were inside ** **


def spec_to_docx(spec_md: str) -> io.BytesIO:
    """Convert the spec Markdown into a simple, clean .docx."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in spec_md.split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif re.match(r"^[-*]\s+", s):
            _add_runs(doc.add_paragraph(style="List Bullet"), re.sub(r"^[-*]\s+", "", s))
        elif re.match(r"^\d+\.\s+", s):
            _add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s+", "", s))
        else:
            _add_runs(doc.add_paragraph(), s)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.post("/export")
async def export(payload: dict):
    """Return the spec as a downloadable Word (.docx) file."""
    spec_md = (payload.get("spec") or "").strip()
    if not spec_md:
        return JSONResponse(status_code=400, content={"ok": False, "error": "empty spec"})
    buf = spec_to_docx(spec_md)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="functional-spec.docx"'},
    )


# --- Prototype builder: hand the spec to the Claude Code CLI ----------------
BUILD_PROMPT = """Read SPEC.md in this directory. Build a single, self-contained
clickable PROTOTYPE in ONE file named index.html.

Requirements:
- Everything inline in index.html: HTML + CSS + JavaScript. No external files,
  no CDNs, no build step, no network calls.
- Use realistic placeholder/mock data so the key screens and flows from the spec
  are demonstrable by clicking around. This is a visual prototype, not a real
  backend.
- Keep it clean and simple — the goal is to show something at the end of a
  meeting, not a finished product.
- Create ONLY index.html. Do not create other files or run any commands."""

BUILD_PROMPT_RESUME = """The functional specification in SPEC.md has been revised.
UPDATE the existing index.html in this directory so it matches the revised spec.
Change only what needs to change to reflect the update — keep the rest of the
prototype intact. Edit index.html in place; do NOT rewrite it from scratch and do
NOT create other files."""

_TOOL_FLAGS = ["--permission-mode", "acceptEdits",
               "--allowed-tools", "Write", "Edit", "Read", "MultiEdit"]


def _safe_id(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]", "", s)[:40]


def _lf_build_worker(job_id: str, project_id: str, spec_md: str, want_resume: bool, out: Path):
    """Background thread: run the Langflow build flow, then serve the returned HTML.

    The BuildPrototype component writes its own copy (and tracks session resume)
    on the Langflow side; here we just persist the returned html so the existing
    /generated static mount can serve it.
    """
    job = BUILD_JOBS[job_id]
    try:
        tweaks = {LANGFLOW_BUILD_NODE: {
            "project_id": project_id, "resume": want_resume,
            "timeout_s": LANGFLOW_BUILD_TIMEOUT,   # raise the component's CLI timeout
        }}
        # HTTP call must outlive the component's own timeout.
        data = _lf_run(LANGFLOW_BUILD_FLOW, spec_md, tweaks, timeout=LANGFLOW_BUILD_TIMEOUT + 60)
        html = data.get("html") or ""
        if data.get("ok") and html and not job.get("cancelled"):
            (out / "index.html").write_text(html, encoding="utf-8")
        job["result"] = data
    except Exception as e:
        job["error"] = str(e)
    finally:
        job["done"] = True


@app.post("/build")
async def build(payload: dict):
    """Kick off a headless Claude Code run that turns the spec into a prototype.

    Auth is via the CLI's own credentials — CLAUDE_CODE_OAUTH_TOKEN in the
    environment (a Claude Code subscription token), inherited by the subprocess.
    With resume=true and a known session, the CLI edits the existing index.html
    instead of regenerating (faster, fewer tokens).
    """
    spec_md = (payload.get("spec") or "").strip()
    if not spec_md:
        return JSONResponse(status_code=400, content={"ok": False, "error": "empty spec"})

    project_id = _safe_id(str(payload.get("project_id") or "")) or uuid.uuid4().hex[:8]
    want_resume = bool(payload.get("resume"))

    # Langflow path: run the build flow in a background thread (it blocks ~30-90s),
    # keeping the same job_id + polling contract the browser already uses.
    if USE_LANGFLOW:
        out = GEN_DIR / project_id
        out.mkdir(parents=True, exist_ok=True)
        (out / "SPEC.md").write_text(spec_md, encoding="utf-8")
        job_id = uuid.uuid4().hex[:8]
        job = {"mode": "langflow", "dir": out, "project_id": project_id,
               "started": time.time(), "done": False, "result": None,
               "error": None, "cancelled": False}
        BUILD_JOBS[job_id] = job
        t = threading.Thread(target=_lf_build_worker,
                             args=(job_id, project_id, spec_md, want_resume, out), daemon=True)
        job["thread"] = t
        t.start()
        return {"ok": True, "job_id": job_id, "project_id": project_id, "resumed": False}

    proj = PROJECTS.get(project_id)

    out = GEN_DIR / project_id
    out.mkdir(parents=True, exist_ok=True)
    (out / "SPEC.md").write_text(spec_md, encoding="utf-8")
    log_path, err_path = out / "build.log", out / "build.err"
    index_path, backup_path = out / "index.html", out / ".index.backup.html"

    session_id = proj["session_id"] if proj else None
    do_resume = bool(want_resume and session_id and index_path.exists())

    backup = None
    if do_resume:
        # Protect the last-good prototype: restored if this run fails or is cancelled.
        try:
            shutil.copyfile(index_path, backup_path); backup = backup_path
        except Exception:
            backup = None
        cmd = [CLAUDE_BIN, "--resume", session_id, "-p", BUILD_PROMPT_RESUME, "--output-format", "json", *_TOOL_FLAGS]
    else:
        cmd = [CLAUDE_BIN, "-p", BUILD_PROMPT, "--output-format", "json", *_TOOL_FLAGS]

    job_id = uuid.uuid4().hex[:8]
    try:
        proc = subprocess.Popen(
            cmd, cwd=str(out),
            stdout=open(log_path, "w", encoding="utf-8"),   # JSON result (usage/session_id)
            stderr=open(err_path, "w", encoding="utf-8"),
            stdin=subprocess.DEVNULL,
        )
    except FileNotFoundError:
        return JSONResponse(status_code=500, content={
            "ok": False, "error": "Claude Code CLI ('claude') not found on PATH."})
    except Exception as e:
        return JSONResponse(status_code=500, content={"ok": False, "error": str(e)})

    PROJECTS.setdefault(project_id, {"dir": out, "session_id": session_id})
    PROJECTS[project_id]["dir"] = out
    BUILD_JOBS[job_id] = {
        "proc": proc, "dir": out, "log": log_path, "err": err_path, "started": time.time(),
        "project_id": project_id, "backup": backup, "finalized": False,
    }
    return {"ok": True, "job_id": job_id, "project_id": project_id, "resumed": do_resume}


def _build_tokens(stdout_text: str):
    """Pull Claude Code token usage from its --output-format json result."""
    try:
        data = json.loads(stdout_text)
    except Exception:
        return None
    u = data.get("usage") or {}
    inp = (u.get("input_tokens", 0) or 0) + (u.get("cache_read_input_tokens", 0) or 0) \
        + (u.get("cache_creation_input_tokens", 0) or 0)
    out = u.get("output_tokens", 0) or 0
    return {"input": inp, "output": out, "total": inp + out, "cost_usd": data.get("total_cost_usd")}


def _finalize_build(job: dict, rc: int, index_ready: bool, stdout_text: str):
    """Run once when a build ends: persist session on success, restore backup on failure."""
    proj = PROJECTS.get(job["project_id"])
    if rc == 0 and index_ready:
        try:
            sid = json.loads(stdout_text).get("session_id")
        except Exception:
            sid = None
        if proj is not None and sid:
            proj["session_id"] = sid           # future rebuilds resume from here
        b = job.get("backup")
        if b and b.exists():
            b.unlink()                          # success — drop the safety copy
    else:
        b = job.get("backup")                   # failed — restore the last-good prototype
        if b and b.exists():
            try:
                shutil.copyfile(b, job["dir"] / "index.html"); b.unlink()
            except Exception:
                pass


@app.get("/build/status/{job_id}")
def build_status(job_id: str):
    """Poll a build job: running / done / error, with the app URL and token usage once ready."""
    job = BUILD_JOBS.get(job_id)
    if not job:
        return JSONResponse(status_code=404, content={"ok": False, "error": "unknown job"})

    # Langflow-backed job: state lives on the background thread, not a subprocess.
    if job.get("mode") == "langflow":
        index_ready = (job["dir"] / "index.html").exists()
        if not job["done"]:
            status = "running"
        elif job["error"]:
            status = "error"
        else:
            data = job["result"] or {}
            status = "done" if (data.get("ok") and index_ready) else \
                     ("done_no_file" if data.get("returncode") == 0 else "error")
        data = job.get("result") or {}
        return {
            "ok": True,
            "status": status,
            "elapsed_s": round(time.time() - job["started"], 1),
            "index_ready": index_ready,
            "app_url": f"/generated/{job['dir'].name}/index.html" if index_ready else None,
            "returncode": data.get("returncode"),
            "tokens": data.get("tokens"),
            "resumed": data.get("resumed"),
            "log": (job.get("error") or data.get("log") or "")[-4000:],
        }

    proc = job["proc"]
    rc = proc.poll()
    index_ready = (job["dir"] / "index.html").exists()

    def _read(p):
        try:
            return p.read_text(encoding="utf-8", errors="replace")
        except Exception:
            return ""
    stdout_text, err_text = _read(job["log"]), _read(job.get("err"))

    tokens = None
    if rc is not None:
        tokens = _build_tokens(stdout_text)
        if not job["finalized"]:
            job["finalized"] = True
            _finalize_build(job, rc, index_ready, stdout_text)
            index_ready = (job["dir"] / "index.html").exists()   # may have been restored
    status = "running" if rc is None else ("done" if (rc == 0 and index_ready)
             else "done_no_file" if rc == 0 else "error")
    return {
        "ok": True,
        "status": status,
        "elapsed_s": round(time.time() - job["started"], 1),
        "index_ready": index_ready,
        "app_url": f"/generated/{job['dir'].name}/index.html" if index_ready else None,
        "returncode": rc,
        "tokens": tokens,
        "log": (err_text or stdout_text)[-4000:],
    }


@app.post("/build/cancel/{job_id}")
def build_cancel(job_id: str):
    """Stop a running build; the partial is discarded and the last-good prototype restored."""
    job = BUILD_JOBS.get(job_id)
    if not job:
        return JSONResponse(status_code=404, content={"ok": False, "error": "unknown job"})

    # Langflow-backed job: we can't interrupt the in-flight HTTP run, but we flag
    # it so the worker won't overwrite the last-good prototype when it returns.
    if job.get("mode") == "langflow":
        job["cancelled"] = True
        return {"ok": True, "cancelled": True}

    proc = job["proc"]
    if proc.poll() is None:
        try:
            proc.terminate()
        except Exception:
            pass
    b = job.get("backup")
    if b and b.exists():
        try:
            shutil.copyfile(b, job["dir"] / "index.html"); b.unlink()
        except Exception:
            pass
    job["finalized"] = True   # don't let status finalize it again (keeps last-completed session)
    return {"ok": True, "cancelled": True}


app.mount("/generated", StaticFiles(directory=GEN_DIR, html=True), name="generated")
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="127.0.0.1", port=8000)
